"""
Script para generar la documentación del sistema Hallazgos ERO en archivos Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Aplica color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level, color_hex="1F3864"):
    """Agrega un encabezado con color personalizado."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return h


def add_info_table(doc, rows, header_color="1F3864", alt_color="D9E1F2"):
    """Crea una tabla de dos columnas (etiqueta | valor)."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(10.5)
        if i == 0:
            set_cell_bg(row.cells[0], header_color)
            set_cell_bg(row.cells[1], header_color)
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif i % 2 == 0:
            set_cell_bg(row.cells[0], alt_color)
            set_cell_bg(row.cells[1], alt_color)
    return table


def add_permission_table(doc, columns, data_rows, header_color="1F3864"):
    """Crea tabla de permisos multi-columna."""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(columns))
    table.style = 'Table Grid'
    # Header
    for j, col in enumerate(columns):
        cell = table.rows[0].cells[j]
        cell.text = col
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row_data in enumerate(data_rows):
        row = table.rows[i + 1]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            if i % 2 == 0:
                set_cell_bg(cell, "EBF3FB")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
            if j > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def page_break(doc):
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1 – MANUAL DE USUARIO
# ═════════════════════════════════════════════════════════════════════════════

def crear_manual_usuario():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ── Portada ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SISTEMA HALLAZGOS ERO")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Manual de Usuario")
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x27, 0x6E, 0xBD)

    doc.add_paragraph()
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = inst.add_run("Fiduprevisora S.A.")
    r3.font.size = Pt(14)
    r3.bold = True

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run(f"Versión 1.0  |  {datetime.date.today().strftime('%d de %B de %Y')}").font.size = Pt(11)

    page_break(doc)

    # ── 1. Introducción ───────────────────────────────────────────────────────
    add_heading(doc, "1. Introducción", 1)
    doc.add_paragraph(
        "El Sistema Hallazgos ERO es una plataforma web corporativa diseñada para la "
        "gestión integral de hallazgos de riesgo operacional. Permite importar, visualizar, "
        "filtrar y dar seguimiento a los hallazgos registrados en los informes de ERO "
        "(Excel), facilitando el control de planes de acción y el cumplimiento de compromisos "
        "dentro de Fiduprevisora S.A."
    )

    add_heading(doc, "1.1 Objetivos del sistema", 2)
    bullets = [
        "Centralizar la información de hallazgos operacionales.",
        "Proporcionar visibilidad en tiempo real del estado de los planes de acción.",
        "Controlar vencimientos y prórrogas de hallazgos.",
        "Gestionar usuarios con control de acceso basado en roles.",
        "Generar métricas y reportes para la toma de decisiones.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "1.2 Acceso al sistema", 2)
    doc.add_paragraph(
        "El sistema está disponible a través del navegador web. Para acceder, ingrese la URL "
        "proporcionada por el área de tecnología y utilice las credenciales asignadas."
    )

    page_break(doc)

    # ── 2. Roles del sistema ──────────────────────────────────────────────────
    add_heading(doc, "2. Roles del Sistema", 1)
    doc.add_paragraph(
        "El sistema cuenta con tres (3) roles jerarquizados. Cada rol define qué información "
        "puede ver el usuario y qué acciones puede ejecutar."
    )

    role_table_data = [
        ("Rol", "Descripción", "Alcance de Datos"),
        ("Vicepresidente", "Nivel más alto de acceso. Administra usuarios y visualiza toda la información.", "Todos los hallazgos de todas las dependencias"),
        ("Directivo",      "Puede cargar archivos Excel, editar hallazgos y gestionar su área.", "Hallazgos de su vicepresidencia"),
        ("Técnico",        "Solo puede consultar y filtrar hallazgos de su área.", "Hallazgos de su vicepresidencia/dependencia"),
    ]

    table = doc.add_table(rows=len(role_table_data), cols=3)
    table.style = 'Table Grid'
    col_widths = [Cm(4), Cm(7), Cm(5.5)]
    for i, row_data in enumerate(role_table_data):
        row = table.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            cell.width = col_widths[j]
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_bg(cell, "1F3864")
            elif i % 2 == 0:
                set_cell_bg(cell, "D9E1F2")

    doc.add_paragraph()

    # Tabla de permisos detallada
    add_heading(doc, "2.1 Matriz de permisos", 2)
    perms_columns = ["Funcionalidad", "Vicepresidente", "Directivo", "Técnico"]
    perms_data = [
        ("Ver todos los hallazgos",          "✔", "✘", "✘"),
        ("Ver hallazgos de su VP/área",       "✔", "✔", "✔"),
        ("Filtrar hallazgos",                 "✔", "✔", "✔"),
        ("Editar hallazgos",                  "✔", "✔", "✘"),
        ("Cargar archivos Excel",             "✔", "✔", "✘"),
        ("Ver historial de cargas",           "✔", "Solo propias", "✘"),
        ("Eliminar carga de Excel",           "✔", "✔", "✘"),
        ("Analizar Excel (sin guardar)",      "✔", "✔", "✔"),
        ("Ver dashboard global",              "✔", "✘", "✘"),
        ("Ver métricas de su área",           "✔", "✔", "✔"),
        ("Crear usuarios",                    "✔", "✘", "✘"),
        ("Editar usuarios",                   "✔", "✘", "✘"),
        ("Desactivar usuarios",               "✔", "✘", "✘"),
        ("Ver listado de usuarios",           "✔", "Solo su área", "✘"),
    ]
    add_permission_table(doc, perms_columns, perms_data)

    page_break(doc)

    # ── 3. Inicio de sesión ───────────────────────────────────────────────────
    add_heading(doc, "3. Inicio de Sesión", 1)
    doc.add_paragraph(
        "Al ingresar a la URL del sistema, se muestra la pantalla de inicio de sesión."
    )
    add_heading(doc, "Pasos para iniciar sesión:", 2)
    steps = [
        "Ingrese su correo electrónico corporativo en el campo 'Correo electrónico'.",
        "Ingrese su contraseña en el campo 'Contraseña'.",
        "Haga clic en el botón 'Iniciar sesión'.",
        "Si las credenciales son correctas, será redirigido al módulo principal.",
        "Si los datos son incorrectos, aparecerá un mensaje de error.",
    ]
    for s in steps:
        p = doc.add_paragraph(s, style='List Number')
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    add_heading(doc, "Notas de seguridad:", 2)
    notes = [
        "La sesión expira automáticamente después de 8 horas de inactividad.",
        "No comparta su contraseña con otros usuarios.",
        "Para recuperar su contraseña, contacte al administrador del sistema.",
    ]
    for n in notes:
        p = doc.add_paragraph(n, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    page_break(doc)

    # ── 4. Módulo: Hallazgos ──────────────────────────────────────────────────
    add_heading(doc, "4. Módulo de Hallazgos", 1)
    doc.add_paragraph(
        "Este módulo permite consultar, filtrar y gestionar todos los hallazgos "
        "cargados en el sistema. Es el módulo central de la plataforma y está disponible "
        "para todos los roles, con información filtrada según el perfil del usuario."
    )

    add_heading(doc, "4.1 Visualización de hallazgos", 2)
    doc.add_paragraph(
        "Al ingresar al módulo, se presenta una lista paginada de hallazgos. Cada registro "
        "muestra la información principal: código de evento, descripción, estado, responsable, "
        "fecha de cierre proyectada y estado del plan de acción."
    )

    add_heading(doc, "4.2 Filtros disponibles", 2)
    filters = [
        ("Estado",                   "Filtra por el estado actual del hallazgo (Abierto, Cerrado, En proceso, etc.)"),
        ("Estado Plan de Acción",    "Filtra por el estado del plan de acción asociado"),
        ("Dependencia",              "Filtra por la dependencia que reporta el hallazgo"),
        ("Vicepresidencia",          "Filtra por vicepresidencia (disponible para todos los roles)"),
        ("Responsable",              "Filtra por el responsable del hallazgo o plan de acción"),
        ("Búsqueda libre",           "Busca por código de evento, descripción o nombre del plan de acción"),
        ("Vencidos",                 "Muestra solo hallazgos con fecha de cierre superada"),
        ("Con prórroga",             "Muestra solo hallazgos que tienen prórroga activa"),
        ("Rango de fecha de cierre", "Filtra hallazgos dentro de un rango de fechas de cierre proyectadas"),
    ]
    table = doc.add_table(rows=1 + len(filters), cols=2)
    table.style = 'Table Grid'
    for j, h in enumerate(["Filtro", "Descripción"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for i, (f, d) in enumerate(filters):
        row = table.rows[i + 1]
        row.cells[0].text = f
        row.cells[1].text = d
        if i % 2 == 0:
            set_cell_bg(row.cells[0], "EBF3FB")
            set_cell_bg(row.cells[1], "EBF3FB")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading(doc, "4.3 Detalle de un hallazgo", 2)
    doc.add_paragraph(
        "Al hacer clic sobre un hallazgo de la lista, se muestra el panel de detalle con "
        "toda la información del evento y sus actividades asociadas."
    )
    detail_fields = [
        "Código de evento", "Descripción", "Fecha inicial y final del evento",
        "Estado", "Observaciones", "Aplicativo afectado",
        "Vicepresidencia y Dependencia",
        "ID del plan de acción", "Nombre y descripción del plan de acción",
        "Estado del plan de acción", "Responsable del plan de acción",
        "Prórroga y fechas de cierre",
        "Actividades vinculadas al hallazgo",
    ]
    for f in detail_fields:
        p = doc.add_paragraph(f, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "4.4 Edición de hallazgos (Directivo y Vicepresidente)", 2)
    doc.add_paragraph(
        "Los usuarios con rol Directivo o Vicepresidente pueden editar los siguientes campos "
        "del hallazgo:"
    )
    editable = [
        "Estado", "Observaciones", "Estado del plan de acción",
        "Responsable del plan de acción", "Estado de la acción",
        "Responsable de la acción", "Prórroga",
        "Fecha de cierre proyectada", "Fecha de cierre final con prórroga",
    ]
    for f in editable:
        p = doc.add_paragraph(f, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Nota: ")
    r.bold = True
    p.add_run(
        "Los campos como código de evento, descripción, fechas del evento, vicepresidencia "
        "y dependencia son de solo lectura y no pueden ser modificados desde la interfaz."
    )

    page_break(doc)

    # ── 5. Módulo: Carga de Excel ─────────────────────────────────────────────
    add_heading(doc, "5. Módulo de Carga de Archivos Excel", 1)
    doc.add_paragraph(
        "Este módulo permite importar hallazgos al sistema desde archivos Excel (.xlsx o .xls) "
        "exportados del sistema ERO. Solo está disponible para usuarios con rol Directivo "
        "o Vicepresidente."
    )

    add_heading(doc, "5.1 Carga de un nuevo archivo", 2)
    upload_steps = [
        "Navegue al módulo 'Cargas' en el menú lateral.",
        "Haga clic en el botón 'Subir Archivo'.",
        "Seleccione el archivo Excel desde su equipo (formato .xlsx o .xls, máximo 16 MB).",
        "El sistema validará y procesará el archivo automáticamente.",
        "Una vez completado, verá un resumen con: registros exitosos, fallidos y errores.",
    ]
    for s in upload_steps:
        p = doc.add_paragraph(s, style='List Number')
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Importante: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(
        "Cada nueva carga de archivo reemplaza TODOS los hallazgos existentes en el sistema. "
        "Asegúrese de que el archivo cargado contenga toda la información actualizada."
    )

    add_heading(doc, "5.2 Análisis previo (simulación)", 2)
    doc.add_paragraph(
        "Antes de guardar, puede analizar el archivo para verificar su contenido sin afectar "
        "los datos del sistema. Use el botón 'Analizar' para obtener una vista previa de los "
        "registros que serían importados."
    )

    add_heading(doc, "5.3 Historial de cargas", 2)
    doc.add_paragraph(
        "El historial muestra todas las cargas realizadas con la siguiente información:"
    )
    history_fields = [
        ("Nombre del archivo",    "Nombre original del archivo cargado"),
        ("Fecha de carga",        "Fecha y hora en que se realizó la carga"),
        ("Usuario",               "Quien realizó la carga"),
        ("Total de registros",    "Cantidad de filas procesadas"),
        ("Exitosos",              "Registros importados correctamente"),
        ("Fallidos",              "Registros con errores de procesamiento"),
        ("Estado",                "Estado de la carga: Procesando, Completado o Error"),
    ]
    add_info_table(doc, [("Campo", "Descripción")] + history_fields)

    add_heading(doc, "5.4 Eliminación de una carga", 2)
    doc.add_paragraph(
        "Para eliminar una carga y todos los hallazgos asociados a ella, haga clic en el "
        "ícono de eliminar en el historial. Se solicitará confirmación antes de proceder. "
        "Esta acción es irreversible."
    )

    add_heading(doc, "5.5 Formato del archivo Excel", 2)
    doc.add_paragraph(
        "El sistema reconoce automáticamente las columnas del archivo ERO mediante patrones "
        "flexibles. Las columnas esperadas incluyen:"
    )
    excel_cols = [
        "Código de evento", "Descripción del evento",
        "Fecha inicial del evento", "Fecha de finalización del evento",
        "Estado del hallazgo", "Vicepresidencia", "Dependencia que reporta",
        "ID del plan de acción", "Nombre del plan de acción",
        "Descripción del plan de acción", "Estado del plan de acción",
        "Responsable del plan de acción", "Estado de la acción",
        "Responsable de la acción", "Prórroga",
        "Fecha de cierre proyectada",
    ]
    for c in excel_cols:
        p = doc.add_paragraph(c, style='List Bullet')
        p.runs[0].font.size = Pt(11)

    page_break(doc)

    # ── 6. Módulo: Dashboard ──────────────────────────────────────────────────
    add_heading(doc, "6. Módulo de Dashboard (Vicepresidente)", 1)
    doc.add_paragraph(
        "El Dashboard está disponible exclusivamente para el rol Vicepresidente y proporciona "
        "una vista ejecutiva con indicadores clave de gestión (KPIs) y gráficas de análisis."
    )

    add_heading(doc, "6.1 Indicadores clave (KPIs)", 2)
    kpis = [
        ("Total de hallazgos",    "Número total de hallazgos cargados en el sistema"),
        ("Hallazgos abiertos",    "Hallazgos con estado activo/abierto"),
        ("Hallazgos cerrados",    "Hallazgos completados o cerrados"),
        ("Cerrados hoy",          "Hallazgos cerrados en la fecha actual"),
        ("Vencidos",              "Hallazgos cuya fecha de cierre proyectada ya pasó"),
        ("Con prórroga",          "Hallazgos que tienen una prórroga activa"),
    ]
    add_info_table(doc, [("Indicador", "Descripción")] + kpis)
    doc.add_paragraph()

    add_heading(doc, "6.2 Gráficas disponibles", 2)
    charts = [
        ("Distribución por estado",               "Muestra la proporción de hallazgos según su estado actual"),
        ("Por dependencia (Top 15)",              "Las 15 dependencias con mayor número de hallazgos"),
        ("Por responsable (Top 15)",              "Los 15 responsables con mayor carga de hallazgos"),
        ("Distribución por estado del plan",      "Estado de los planes de acción (cumplido, en proceso, etc.)"),
        ("Línea de tiempo (12 meses)",            "Evolución mensual de hallazgos abiertos y cerrados"),
        ("Seguimiento de prórrogas",              "Análisis de hallazgos con prórroga activa vs. vencida"),
        ("Cargas recientes",                      "Últimas 5 importaciones de archivos Excel"),
    ]
    add_info_table(doc, [("Gráfica", "Descripción")] + charts)

    page_break(doc)

    # ── 7. Módulo: Gestión de Usuarios ────────────────────────────────────────
    add_heading(doc, "7. Módulo de Gestión de Usuarios (Vicepresidente)", 1)
    doc.add_paragraph(
        "Este módulo permite administrar las cuentas de usuario del sistema. Solo los usuarios "
        "con rol Vicepresidente tienen acceso completo a esta funcionalidad."
    )

    add_heading(doc, "7.1 Creación de usuario", 2)
    create_steps = [
        "Navegue al módulo 'Usuarios' en el menú lateral.",
        "Haga clic en el botón 'Nuevo Usuario'.",
        "Complete los campos requeridos: nombre, correo electrónico, contraseña, rol.",
        "Opcionalmente, asigne vicepresidencia y dependencia.",
        "Haga clic en 'Guardar'.",
    ]
    for s in create_steps:
        p = doc.add_paragraph(s, style='List Number')
        p.runs[0].font.size = Pt(11)

    add_heading(doc, "7.2 Campos del usuario", 2)
    user_fields = [
        ("Nombre completo",    "Nombre del usuario en el sistema — requerido"),
        ("Correo electrónico", "Identificador único de acceso — requerido"),
        ("Contraseña",         "Mínimo 8 caracteres — requerido"),
        ("Rol",                "Vicepresidente / Directivo / Técnico — requerido"),
        ("Vicepresidencia",    "Vicepresidencia a la que pertenece — opcional"),
        ("Dependencia",        "Dependencia o área específica — opcional"),
        ("Estado",             "Activo / Inactivo — por defecto Activo"),
    ]
    add_info_table(doc, [("Campo", "Descripción")] + user_fields)
    doc.add_paragraph()

    add_heading(doc, "7.3 Edición de usuario", 2)
    doc.add_paragraph(
        "Haga clic en el ícono de editar junto al usuario deseado para modificar cualquier "
        "campo. Los cambios se aplican inmediatamente al guardar."
    )

    add_heading(doc, "7.4 Desactivación de usuario", 2)
    doc.add_paragraph(
        "Al desactivar un usuario, este no podrá iniciar sesión, pero su historial se conserva. "
        "El sistema no permite eliminar usuarios permanentemente ni desactivar la propia cuenta."
    )

    page_break(doc)

    # ── 8. Guía por rol ───────────────────────────────────────────────────────
    add_heading(doc, "8. Guía Rápida por Rol", 1)

    # Vicepresidente
    add_heading(doc, "8.1 Vicepresidente", 2)
    vp_tasks = [
        ("Monitorear KPIs globales",         "Dashboard → ver indicadores y gráficas"),
        ("Consultar todos los hallazgos",     "Hallazgos → sin filtro de área"),
        ("Editar cualquier hallazgo",         "Hallazgos → seleccionar → editar"),
        ("Cargar nuevo ERO",                  "Cargas → Subir Archivo"),
        ("Ver historial completo de cargas",  "Cargas → Historial"),
        ("Crear/editar/desactivar usuarios",  "Usuarios → gestionar"),
    ]
    add_info_table(doc, [("Tarea", "Cómo hacerlo")] + vp_tasks)
    doc.add_paragraph()

    # Directivo
    add_heading(doc, "8.2 Directivo", 2)
    dir_tasks = [
        ("Ver hallazgos de su VP",           "Hallazgos → automáticamente filtrado por su área"),
        ("Editar hallazgo de su área",       "Hallazgos → seleccionar → editar campos permitidos"),
        ("Cargar nuevo ERO",                 "Cargas → Subir Archivo"),
        ("Ver sus cargas",                   "Cargas → Historial (solo sus propias cargas)"),
        ("Consultar métricas de su área",    "Dashboard → métricas filtradas por vicepresidencia"),
    ]
    add_info_table(doc, [("Tarea", "Cómo hacerlo")] + dir_tasks)
    doc.add_paragraph()

    # Técnico
    add_heading(doc, "8.3 Técnico", 2)
    tec_tasks = [
        ("Consultar hallazgos de su área", "Hallazgos → filtrado automático por su dependencia/VP"),
        ("Usar filtros de búsqueda",       "Hallazgos → barra de filtros"),
        ("Ver detalle de un hallazgo",     "Hallazgos → clic en el hallazgo"),
        ("Ver actividades asociadas",      "Detalle del hallazgo → sección Actividades"),
        ("Analizar Excel sin guardar",     "Cargas → Analizar (no guarda datos)"),
        ("Consultar métricas de su área",  "Dashboard → métricas de su vicepresidencia"),
    ]
    add_info_table(doc, [("Tarea", "Cómo hacerlo")] + tec_tasks)

    page_break(doc)

    # ── 9. Preguntas frecuentes ───────────────────────────────────────────────
    add_heading(doc, "9. Preguntas Frecuentes", 1)
    faqs = [
        (
            "¿Por qué no puedo ver todos los hallazgos?",
            "El sistema filtra automáticamente los hallazgos según su rol y área asignada. "
            "Solo el Vicepresidente puede ver todos los hallazgos."
        ),
        (
            "¿Qué ocurre si cargo un nuevo archivo Excel?",
            "Todos los hallazgos existentes serán reemplazados por los del nuevo archivo. "
            "Esta acción no se puede deshacer, por lo que se recomienda verificar el archivo "
            "con la opción de análisis previo antes de cargar."
        ),
        (
            "¿Puedo editar un hallazgo si soy Técnico?",
            "No. El rol Técnico es de solo consulta. Solo Directivos y Vicepresidentes "
            "pueden editar información de hallazgos."
        ),
        (
            "¿La sesión expira?",
            "Sí, la sesión expira automáticamente después de 8 horas. Deberá iniciar sesión "
            "nuevamente para continuar."
        ),
        (
            "¿Cómo puedo cambiar mi contraseña?",
            "Contacte al administrador del sistema (usuario con rol Vicepresidente) para que "
            "actualice su contraseña."
        ),
        (
            "¿Qué formato debe tener el archivo Excel?",
            "El archivo debe ser .xlsx o .xls, exportado directamente del sistema ERO. "
            "El tamaño máximo es 16 MB. Las columnas son detectadas automáticamente."
        ),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run(f"P: {q}")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        doc.add_paragraph(f"R: {a}").runs[0].font.size = Pt(11)
        doc.add_paragraph()

    path = "c:/Users/Usuario/hallazgos/Manual_de_Usuario.docx"
    doc.save(path)
    print(f"[OK] Guardado: {path}")
    return path


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2 – DIAGRAMA DE RELACIONES
# ═════════════════════════════════════════════════════════════════════════════

def crear_diagrama_relaciones():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Portada
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SISTEMA HALLAZGOS ERO")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Diagrama de Relaciones y Modelo de Datos")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x27, 0x6E, 0xBD)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run(f"Versión 1.0  |  {datetime.date.today().strftime('%d de %B de %Y')}").font.size = Pt(11)

    page_break(doc)

    # ── 1. Descripción general ────────────────────────────────────────────────
    add_heading(doc, "1. Descripción General del Modelo", 1)
    doc.add_paragraph(
        "El sistema Hallazgos ERO utiliza una base de datos PostgreSQL con cuatro entidades "
        "principales relacionadas entre sí. A continuación se detalla cada entidad, sus atributos "
        "y las relaciones que las conectan."
    )

    # ── 2. Diagrama textual ───────────────────────────────────────────────────
    add_heading(doc, "2. Diagrama de Relaciones (Modelo Entidad-Relación)", 1)

    er_diagram = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                        MODELO ENTIDAD-RELACIÓN                              │
│                        Sistema Hallazgos ERO                                │
└─────────────────────────────────────────────────────────────────────────────┘

  ┌──────────────────┐          ┌────────────────────────┐
  │      USUARIO     │  1      N │    UPLOAD_HISTORY      │
  │──────────────────│──────────▶│────────────────────────│
  │ PK id            │          │ PK id                  │
  │    nombre        │          │    filename_original   │
  │    email (UNIQUE)│          │    filename_stored     │
  │    password_hash │          │ FK uploaded_by_id──────│── → USUARIO.id
  │    rol (ENUM)    │          │    total_registros     │
  │    vicepresidencia          │    registros_exitosos  │
  │    dependencia   │          │    registros_fallidos  │
  │    activo (BOOL) │          │    estado (ENUM)       │
  │    created_at    │          │    errores (TEXT)      │
  │    updated_at    │          │    uploaded_at         │
  └──────────────────┘          └────────────────────────┘
                                          │ 1
                                          │
                                          │ N
                               ┌──────────▼───────────────────────────────────┐
                               │               HALLAZGO                       │
                               │──────────────────────────────────────────────│
                               │ PK id                                        │
                               │    codigo_evento (INDEX)                     │
                               │    descripcion (TEXT)                        │
                               │    fecha_inicial_evento                      │
                               │    fecha_finalizacion_evento                 │
                               │    estado (INDEX)                            │
                               │    observaciones (TEXT)                      │
                               │    aplicativo_afecta_ero                     │
                               │    vicepresidencia (INDEX)                   │
                               │    dependencia_reporta_ero (INDEX)           │
                               │    reportado_para                            │
                               │    reportado_por                             │
                               │    id_plan_accion (INDEX)                    │
                               │    nombre_plan_accion                        │
                               │    descripcion_plan_accion (TEXT)            │
                               │    estado_plan_accion (INDEX)                │
                               │    responsable_plan_accion (INDEX)           │
                               │    estado_accion                             │
                               │    responsable_accion (INDEX)                │
                               │    prorroga                                  │
                               │    fecha_cierre_proyectada                   │
                               │    fecha_cierre_final_prorroga               │
                               │ FK upload_id (INDEX) ────────────────────────│── → UPLOAD_HISTORY.id
                               │    created_at                                │
                               │    updated_at                                │
                               └──────────────────────────────────────────────┘
                                          │ 1
                                          │  (CASCADE DELETE)
                                          │ N
                               ┌──────────▼───────────────────────────────────┐
                               │               ACTIVIDAD                      │
                               │──────────────────────────────────────────────│
                               │ PK id                                        │
                               │ FK hallazgo_id (INDEX) ──────────────────────│── → HALLAZGO.id
                               │    codigo_evento (INDEX)                     │
                               │    id_plan_accion                            │
                               │    nombre_plan_accion                        │
                               │    descripcion (TEXT)                        │
                               │    estado_plan_accion                        │
                               │    responsable                               │
                               │    estado_accion                             │
                               │    responsable_accion                        │
                               │    fecha_compromiso                          │
                               │    prorroga                                  │
                               │    fecha_prorroga                            │
                               │    observaciones (TEXT)                      │
                               │    created_at                                │
                               └──────────────────────────────────────────────┘

  LEYENDA:
  PK = Clave Primaria      FK = Clave Foránea
  INDEX = Campo indexado   UNIQUE = Valor único
  ENUM = Tipo enumerado    1..N = Relación uno a muchos
  CASCADE DELETE = Al eliminar el padre, se eliminan los hijos automáticamente
"""

    p = doc.add_paragraph()
    run = p.add_run(er_diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    page_break(doc)

    # ── 3. Descripción de entidades ───────────────────────────────────────────
    add_heading(doc, "3. Descripción de Entidades", 1)

    # USUARIO
    add_heading(doc, "3.1 Entidad: USUARIO", 2)
    doc.add_paragraph(
        "Almacena los usuarios que tienen acceso al sistema. Cada usuario tiene un rol "
        "que determina sus permisos."
    )
    user_attrs = [
        ("id",             "INTEGER, PK",     "Identificador único autoincremental"),
        ("nombre",         "VARCHAR",          "Nombre completo del usuario"),
        ("email",          "VARCHAR, UNIQUE",  "Correo electrónico — usado para login"),
        ("password_hash",  "VARCHAR",          "Contraseña cifrada con bcrypt"),
        ("rol",            "ENUM",             "vicepresidente | directivo | tecnico"),
        ("vicepresidencia","VARCHAR, NULL",     "VP a la que pertenece el usuario"),
        ("dependencia",    "VARCHAR, NULL",     "Dependencia/área del usuario"),
        ("activo",         "BOOLEAN",           "True=activo | False=desactivado"),
        ("created_at",     "DATETIME",          "Fecha de creación del registro"),
        ("updated_at",     "DATETIME",          "Última modificación del registro"),
    ]
    t = doc.add_table(rows=1 + len(user_attrs), cols=3)
    t.style = 'Table Grid'
    for j, h in enumerate(["Atributo", "Tipo", "Descripción"]):
        cell = t.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for i, (a, tp, d) in enumerate(user_attrs):
        row = t.rows[i + 1]
        row.cells[0].text = a
        row.cells[1].text = tp
        row.cells[2].text = d
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, "EBF3FB")
        for c in row.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # HALLAZGO
    add_heading(doc, "3.2 Entidad: HALLAZGO", 2)
    doc.add_paragraph(
        "Entidad principal del sistema. Representa un hallazgo de riesgo operacional importado "
        "desde el informe ERO. Cada hallazgo agrupa un código de evento con su plan de acción."
    )
    hallazgo_attrs = [
        ("id",                          "INTEGER, PK",    "Identificador único"),
        ("codigo_evento",               "VARCHAR, INDEX", "Código único del evento ERO"),
        ("descripcion",                 "TEXT",           "Descripción detallada del hallazgo"),
        ("fecha_inicial_evento",        "DATETIME",       "Fecha de inicio del evento"),
        ("fecha_finalizacion_evento",   "DATETIME",       "Fecha de fin del evento"),
        ("estado",                      "VARCHAR, INDEX", "Estado actual del hallazgo"),
        ("observaciones",               "TEXT",           "Notas y observaciones editables"),
        ("aplicativo_afecta_ero",       "VARCHAR",        "Sistema o aplicativo afectado"),
        ("vicepresidencia",             "VARCHAR, INDEX", "VP que reporta el hallazgo"),
        ("dependencia_reporta_ero",     "VARCHAR, INDEX", "Dependencia que reporta"),
        ("reportado_para",              "VARCHAR",        "Destinatario del reporte"),
        ("reportado_por",               "VARCHAR",        "Responsable del reporte"),
        ("id_plan_accion",              "VARCHAR, INDEX", "ID del plan de acción"),
        ("nombre_plan_accion",          "VARCHAR",        "Nombre descriptivo del plan"),
        ("descripcion_plan_accion",     "TEXT",           "Descripción del plan de acción"),
        ("estado_plan_accion",          "VARCHAR, INDEX", "Estado del plan de acción"),
        ("responsable_plan_accion",     "VARCHAR, INDEX", "Responsable del plan de acción"),
        ("estado_accion",               "VARCHAR",        "Estado de la acción específica"),
        ("responsable_accion",          "VARCHAR, INDEX", "Responsable de la acción"),
        ("prorroga",                    "VARCHAR",        "Indicador de prórroga"),
        ("fecha_cierre_proyectada",     "DATETIME",       "Fecha de cierre planificada"),
        ("fecha_cierre_final_prorroga", "DATETIME",       "Fecha final con prórroga"),
        ("upload_id",                   "INTEGER, FK",    "Referencia a la carga de origen"),
        ("created_at",                  "DATETIME",       "Fecha de creación"),
        ("updated_at",                  "DATETIME",       "Última modificación"),
    ]
    t2 = doc.add_table(rows=1 + len(hallazgo_attrs), cols=3)
    t2.style = 'Table Grid'
    for j, h in enumerate(["Atributo", "Tipo", "Descripción"]):
        cell = t2.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for i, (a, tp, d) in enumerate(hallazgo_attrs):
        row = t2.rows[i + 1]
        row.cells[0].text = a
        row.cells[1].text = tp
        row.cells[2].text = d
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, "EBF3FB")
        for c in row.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # ACTIVIDAD
    add_heading(doc, "3.3 Entidad: ACTIVIDAD", 2)
    doc.add_paragraph(
        "Representa las actividades del plan de acción vinculadas a un hallazgo. Un hallazgo "
        "puede tener múltiples actividades. Si al eliminar el hallazgo, las actividades se "
        "eliminan en cascada."
    )
    act_attrs = [
        ("id",                  "INTEGER, PK",    "Identificador único"),
        ("hallazgo_id",         "INTEGER, FK",    "Referencia al hallazgo padre (CASCADE)"),
        ("codigo_evento",       "VARCHAR, INDEX", "Código del evento para referencia"),
        ("id_plan_accion",      "VARCHAR",        "ID del plan de acción"),
        ("nombre_plan_accion",  "VARCHAR",        "Nombre del plan de acción"),
        ("descripcion",         "TEXT",           "Descripción de la actividad"),
        ("estado_plan_accion",  "VARCHAR",        "Estado del plan"),
        ("responsable",         "VARCHAR",        "Responsable de la actividad"),
        ("estado_accion",       "VARCHAR",        "Estado específico de la acción"),
        ("responsable_accion",  "VARCHAR",        "Responsable de la acción"),
        ("fecha_compromiso",    "DATETIME",       "Fecha comprometida de ejecución"),
        ("prorroga",            "VARCHAR",        "Indicador de prórroga"),
        ("fecha_prorroga",      "DATETIME",       "Fecha de prórroga si aplica"),
        ("observaciones",       "TEXT",           "Notas adicionales"),
        ("created_at",          "DATETIME",       "Fecha de creación del registro"),
    ]
    t3 = doc.add_table(rows=1 + len(act_attrs), cols=3)
    t3.style = 'Table Grid'
    for j, h in enumerate(["Atributo", "Tipo", "Descripción"]):
        cell = t3.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for i, (a, tp, d) in enumerate(act_attrs):
        row = t3.rows[i + 1]
        row.cells[0].text = a
        row.cells[1].text = tp
        row.cells[2].text = d
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, "EBF3FB")
        for c in row.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # UPLOAD_HISTORY
    add_heading(doc, "3.4 Entidad: UPLOAD_HISTORY", 2)
    doc.add_paragraph(
        "Registra el historial de cargas de archivos Excel para auditoría y trazabilidad. "
        "Cada carga está vinculada al usuario que la realizó y puede contener múltiples hallazgos."
    )
    upload_attrs = [
        ("id",                   "INTEGER, PK",  "Identificador único"),
        ("filename_original",    "VARCHAR",       "Nombre original del archivo"),
        ("filename_stored",      "VARCHAR",       "Nombre del archivo almacenado"),
        ("uploaded_by_id",       "INTEGER, FK",  "Usuario que realizó la carga"),
        ("total_registros",      "INTEGER",       "Total de filas procesadas"),
        ("registros_exitosos",   "INTEGER",       "Filas importadas correctamente"),
        ("registros_fallidos",   "INTEGER",       "Filas con error de procesamiento"),
        ("estado",               "ENUM",          "procesando | completado | error"),
        ("errores",              "TEXT",          "Detalle de errores encontrados"),
        ("uploaded_at",          "DATETIME",      "Fecha y hora de la carga"),
    ]
    t4 = doc.add_table(rows=1 + len(upload_attrs), cols=3)
    t4.style = 'Table Grid'
    for j, h in enumerate(["Atributo", "Tipo", "Descripción"]):
        cell = t4.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F3864")
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for i, (a, tp, d) in enumerate(upload_attrs):
        row = t4.rows[i + 1]
        row.cells[0].text = a
        row.cells[1].text = tp
        row.cells[2].text = d
        if i % 2 == 0:
            for c in row.cells:
                set_cell_bg(c, "EBF3FB")
        for c in row.cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    page_break(doc)

    # ── 4. Relaciones ─────────────────────────────────────────────────────────
    add_heading(doc, "4. Relaciones entre Entidades", 1)
    relations = [
        (
            "USUARIO → UPLOAD_HISTORY",
            "Uno a Muchos (1:N)",
            "Un usuario puede realizar múltiples cargas de archivos. "
            "La FK uploaded_by_id en UPLOAD_HISTORY apunta a USUARIO.id."
        ),
        (
            "UPLOAD_HISTORY → HALLAZGO",
            "Uno a Muchos (1:N)",
            "Una carga puede generar múltiples hallazgos. "
            "La FK upload_id en HALLAZGO apunta a UPLOAD_HISTORY.id."
        ),
        (
            "HALLAZGO → ACTIVIDAD",
            "Uno a Muchos (1:N) con CASCADE DELETE",
            "Un hallazgo puede tener múltiples actividades de plan de acción. "
            "Al eliminar un hallazgo, todas sus actividades se eliminan automáticamente. "
            "La FK hallazgo_id en ACTIVIDAD apunta a HALLAZGO.id."
        ),
    ]
    for rel, tipo, desc in relations:
        p = doc.add_paragraph()
        r = p.add_run(rel)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"Tipo: {tipo}")
        r2.bold = True
        r2.font.size = Pt(10)
        doc.add_paragraph(desc).runs[0].font.size = Pt(10)
        doc.add_paragraph()

    page_break(doc)

    # ── 5. Arquitectura del sistema ───────────────────────────────────────────
    add_heading(doc, "5. Arquitectura del Sistema", 1)

    arch_diagram = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                         ARQUITECTURA DEL SISTEMA                            │
└─────────────────────────────────────────────────────────────────────────────┘

  ┌──────────────────────────────────┐
  │         USUARIO FINAL            │
  │    (Navegador web / Browser)     │
  └──────────────────┬───────────────┘
                     │ HTTP/HTTPS
                     ▼
  ┌──────────────────────────────────┐
  │       FRONTEND                   │
  │  Next.js 16 + React 19           │
  │  TypeScript + Tailwind CSS       │
  │  shadcn/ui + Recharts            │
  │  Puerto: 3000                    │
  └──────────────────┬───────────────┘
                     │ REST API / JSON
                     │ Authorization: Bearer <JWT>
                     ▼
  ┌──────────────────────────────────┐
  │       BACKEND (API REST)         │
  │  Flask 3 + Python 3.11           │
  │  Flask-JWT-Extended              │
  │  Flask-SQLAlchemy                │
  │  Gunicorn (producción)           │
  │  Puerto: 5000                    │
  └──────────────────┬───────────────┘
                     │ SQLAlchemy ORM
                     ▼
  ┌──────────────────────────────────┐
  │       BASE DE DATOS              │
  │  PostgreSQL 15                   │
  │  4 tablas: usuario,              │
  │  hallazgo, actividad,            │
  │  upload_history                  │
  │  Puerto: 5432                    │
  └──────────────────────────────────┘

  ┌──────────────────────────────────┐
  │       ORQUESTACIÓN               │
  │  Docker Compose v3.9             │
  │  3 contenedores:                 │
  │  - postgres                      │
  │  - backend                       │
  │  - frontend                      │
  └──────────────────────────────────┘

  SEGURIDAD:
  ┌─────────────────────────────────────────────────────────────────────┐
  │  • JWT (HS256) con expiración de 8 horas                            │
  │  • Contraseñas cifradas con bcrypt (werkzeug.security)              │
  │  • RBAC (Control de acceso basado en roles) en backend              │
  │  • Filtrado de datos a nivel de consulta SQL según rol              │
  │  • CORS configurado para orígenes específicos                       │
  └─────────────────────────────────────────────────────────────────────┘
"""

    p = doc.add_paragraph()
    run = p.add_run(arch_diagram)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    page_break(doc)

    # ── 6. Flujo de datos ─────────────────────────────────────────────────────
    add_heading(doc, "6. Flujo de Datos – Carga de Excel", 1)

    flow = """
  Archivo .xlsx/.xls
        │
        ▼
  [FRONTEND] Validación cliente
  • Extensión: .xlsx o .xls
  • Tamaño máximo: 16 MB
        │
        ▼
  [BACKEND] Recepción y almacenamiento temporal
  • Guarda archivo con nombre UUID-prefixed
  • Crea registro UPLOAD_HISTORY (estado=procesando)
        │
        ▼
  [EXCEL PARSER] excel_parser.py
  • Detecta columnas por regex (flexible)
  • Agrupa filas por codigo_evento
  • Primera fila del grupo → HALLAZGO
  • Filas adicionales del grupo → ACTIVIDAD
  • Parsea fechas (7 formatos soportados)
        │
        ├── Éxito parcial: guarda hallazgos/actividades válidos
        │   registra filas con error
        │
        └── Transacción exitosa:
            • Elimina hallazgos anteriores (carga global)
            • Inserta nuevos HALLAZGO + ACTIVIDAD
            • Actualiza UPLOAD_HISTORY (estado=completado)
        │
        ▼
  [FRONTEND] Muestra resumen:
  • Registros exitosos
  • Registros fallidos
  • Lista de errores (máximo 10)
"""

    p = doc.add_paragraph()
    run = p.add_run(flow)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    path = "c:/Users/Usuario/hallazgos/Diagrama_Relaciones.docx"
    doc.save(path)
    print(f"[OK] Guardado: {path}")
    return path


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 3 – HISTORIAS DE USUARIO
# ═════════════════════════════════════════════════════════════════════════════

def crear_historias_usuario():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Portada
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SISTEMA HALLAZGOS ERO")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Historias de Usuario")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x27, 0x6E, 0xBD)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run(f"Versión 1.0  |  {datetime.date.today().strftime('%d de %B de %Y')}").font.size = Pt(11)

    page_break(doc)

    # Introducción
    add_heading(doc, "Introducción", 1)
    doc.add_paragraph(
        "Este documento contiene las historias de usuario del Sistema Hallazgos ERO, "
        "organizadas por épica funcional. Cada historia describe una necesidad concreta "
        "del usuario, los criterios de aceptación y la prioridad de implementación."
    )

    add_heading(doc, "Convenciones", 2)
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    convs = [
        ("Prioridad", "Alta / Media / Baja"),
        ("Rol del usuario", "Vicepresidente / Directivo / Técnico"),
        ("Identificador", "HU-XX donde XX es número secuencial"),
        ("Criterios de aceptación", "Condiciones verificables que confirman la implementación"),
    ]
    set_cell_bg(t.rows[0].cells[0], "1F3864")
    set_cell_bg(t.rows[0].cells[1], "1F3864")
    for i, (k, v) in enumerate(convs):
        row = t.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if i % 2 == 0 and i > 0:
            set_cell_bg(row.cells[0], "EBF3FB")
            set_cell_bg(row.cells[1], "EBF3FB")

    page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # Helper para agregar historia de usuario
    # ─────────────────────────────────────────────────────────────────────────
    def add_story(doc, hu_id, title, role, story, criteria, priority, epic):
        add_heading(doc, f"{hu_id}: {title}", 2)

        meta_rows = [
            ("ID",        hu_id),
            ("Épica",     epic),
            ("Prioridad", priority),
            ("Rol",       role),
        ]
        t = doc.add_table(rows=len(meta_rows), cols=2)
        t.style = 'Table Grid'
        for i, (k, v) in enumerate(meta_rows):
            t.rows[i].cells[0].text = k
            t.rows[i].cells[1].text = v
            t.rows[i].cells[0].width = Cm(3)
            t.rows[i].cells[1].width = Cm(13)
            if i == 0:
                set_cell_bg(t.rows[i].cells[0], "1F3864")
                set_cell_bg(t.rows[i].cells[1], "1F3864")
                for cell in t.rows[i].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
            else:
                if i % 2 == 0:
                    set_cell_bg(t.rows[i].cells[0], "D9E1F2")
                    set_cell_bg(t.rows[i].cells[1], "D9E1F2")
                for cell in t.rows[i].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()

        # Historia
        p = doc.add_paragraph()
        r = p.add_run("Historia de usuario:")
        r.bold = True
        r.font.size = Pt(11)
        doc.add_paragraph(
            f"Como {role}, quiero {story}."
        ).runs[0].font.size = Pt(11)

        # Criterios
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Criterios de aceptación:")
        r2.bold = True
        r2.font.size = Pt(11)
        for c in criteria:
            p3 = doc.add_paragraph(c, style='List Bullet')
            p3.runs[0].font.size = Pt(11)

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 1: AUTENTICACIÓN
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 1: Autenticación y Seguridad", 1)

    add_story(doc,
        "HU-01", "Inicio de sesión con credenciales",
        "usuario del sistema",
        "poder iniciar sesión con mi correo y contraseña para acceder a las funcionalidades de mi rol",
        [
            "El sistema muestra un formulario con campos de correo y contraseña.",
            "Al ingresar credenciales correctas, se redirige al módulo principal.",
            "Al ingresar credenciales incorrectas, se muestra mensaje de error descriptivo.",
            "Las contraseñas no se muestran en texto plano.",
            "La sesión persiste durante 8 horas sin necesidad de volver a iniciar sesión.",
            "Solo usuarios con estado 'activo' pueden iniciar sesión.",
        ],
        "Alta", "Autenticación"
    )

    add_story(doc,
        "HU-02", "Cierre de sesión",
        "usuario autenticado",
        "poder cerrar mi sesión de forma segura para proteger mis datos",
        [
            "Existe un botón/opción de cerrar sesión visible en la interfaz.",
            "Al cerrar sesión, se eliminan los datos de sesión del navegador.",
            "Tras cerrar sesión, se redirige a la pantalla de inicio de sesión.",
            "El token JWT deja de ser válido para el cliente.",
        ],
        "Alta", "Autenticación"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 2: GESTIÓN DE HALLAZGOS
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 2: Gestión de Hallazgos", 1)

    add_story(doc,
        "HU-03", "Visualizar listado de hallazgos",
        "usuario autenticado",
        "ver una lista paginada de hallazgos para consultar el estado de los eventos de riesgo operacional",
        [
            "Se muestran mínimo los siguientes campos: código de evento, descripción, estado, responsable, fecha de cierre, estado del plan.",
            "La lista está paginada (por defecto 20 registros por página).",
            "El Vicepresidente ve todos los hallazgos.",
            "El Directivo ve solo los hallazgos de su vicepresidencia.",
            "El Técnico ve solo los hallazgos de su vicepresidencia o dependencia.",
        ],
        "Alta", "Hallazgos"
    )

    add_story(doc,
        "HU-04", "Filtrar hallazgos por criterios múltiples",
        "usuario autenticado",
        "filtrar los hallazgos por estado, responsable, área, fechas y otros criterios para encontrar rápidamente los registros relevantes",
        [
            "Se puede filtrar por estado del hallazgo.",
            "Se puede filtrar por estado del plan de acción.",
            "Se puede filtrar por vicepresidencia y dependencia.",
            "Se puede filtrar por responsable.",
            "Se puede realizar búsqueda de texto libre (código, descripción, nombre del plan).",
            "Se puede filtrar hallazgos vencidos (fecha de cierre superada).",
            "Se puede filtrar hallazgos con prórroga.",
            "Se puede filtrar por rango de fechas de cierre.",
            "Los filtros se pueden combinar.",
            "Al limpiar filtros, se muestra el listado completo.",
        ],
        "Alta", "Hallazgos"
    )

    add_story(doc,
        "HU-05", "Ver detalle de un hallazgo",
        "usuario autenticado",
        "ver el detalle completo de un hallazgo incluyendo sus actividades para entender el contexto y estado del evento",
        [
            "Al seleccionar un hallazgo, se muestra toda su información.",
            "Se muestran las actividades vinculadas al hallazgo.",
            "Los campos de solo lectura no son editables.",
            "Se indica claramente si el hallazgo está vencido.",
        ],
        "Alta", "Hallazgos"
    )

    add_story(doc,
        "HU-06", "Editar campos de un hallazgo",
        "Directivo o Vicepresidente",
        "editar los campos permitidos de un hallazgo para actualizar el estado del plan de acción y observaciones",
        [
            "Solo los roles Directivo y Vicepresidente pueden editar hallazgos.",
            "Los campos editables son: estado, observaciones, estado del plan de acción, responsable del plan, estado de la acción, responsable de la acción, prórroga, fechas de cierre.",
            "Los cambios se guardan al confirmar la edición.",
            "Se muestra confirmación visual al guardar exitosamente.",
            "Los campos de solo lectura (código, descripción, fechas del evento, vicepresidencia, dependencia) no son modificables.",
        ],
        "Alta", "Hallazgos"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 3: CARGA DE ARCHIVOS EXCEL
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 3: Carga de Archivos Excel", 1)

    add_story(doc,
        "HU-07", "Importar hallazgos desde Excel",
        "Directivo o Vicepresidente",
        "cargar un archivo Excel exportado del sistema ERO para actualizar los hallazgos en el sistema",
        [
            "Solo los roles Directivo y Vicepresidente pueden cargar archivos.",
            "Se aceptan formatos .xlsx y .xls con tamaño máximo de 16 MB.",
            "El sistema detecta automáticamente las columnas del ERO.",
            "Al cargar, se reemplazan todos los hallazgos previos.",
            "Al finalizar, se muestra un resumen: registros exitosos, fallidos y lista de errores.",
            "Si el archivo es inválido, se muestra un mensaje de error descriptivo.",
        ],
        "Alta", "Carga Excel"
    )

    add_story(doc,
        "HU-08", "Analizar Excel antes de cargar",
        "usuario autenticado",
        "analizar el contenido de un archivo Excel sin guardar los datos para verificar que el formato es correcto antes de reemplazar los hallazgos existentes",
        [
            "Existe opción de analizar sin guardar.",
            "El análisis muestra cuántos registros serían importados.",
            "Se indica si hay filas con errores de formato.",
            "No se modifica ningún dato en la base de datos durante el análisis.",
        ],
        "Media", "Carga Excel"
    )

    add_story(doc,
        "HU-09", "Ver historial de cargas",
        "Directivo o Vicepresidente",
        "ver el historial de cargas realizadas para realizar auditoría y control de las importaciones",
        [
            "El Vicepresidente ve todas las cargas de todos los usuarios.",
            "El Directivo ve solo sus propias cargas.",
            "Cada registro muestra: nombre del archivo, fecha, usuario, totales y estado.",
            "Se puede ver el detalle de errores de cada carga.",
            "La lista está paginada.",
        ],
        "Media", "Carga Excel"
    )

    add_story(doc,
        "HU-10", "Eliminar una carga y sus hallazgos",
        "Directivo o Vicepresidente",
        "eliminar una carga del historial junto con todos sus hallazgos asociados cuando la información es incorrecta o desactualizada",
        [
            "Se solicita confirmación antes de eliminar.",
            "Al eliminar, se borran todos los hallazgos y actividades de esa carga.",
            "Se muestra mensaje de confirmación tras la eliminación exitosa.",
            "La eliminación es irreversible.",
        ],
        "Media", "Carga Excel"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 4: DASHBOARD
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 4: Dashboard y Métricas", 1)

    add_story(doc,
        "HU-11", "Ver dashboard ejecutivo global",
        "Vicepresidente",
        "ver un dashboard con indicadores clave globales para tomar decisiones basadas en datos actualizados de toda la organización",
        [
            "El dashboard muestra: total de hallazgos, abiertos, cerrados, cerrados hoy, vencidos y con prórroga.",
            "Se muestran gráficas: distribución por estado, por dependencia, por responsable, por estado del plan.",
            "Se muestra una línea de tiempo de los últimos 12 meses.",
            "Se muestra seguimiento de prórrogas.",
            "Los datos son actualizados en tiempo real al cargar la página.",
            "Solo el Vicepresidente accede a esta vista global.",
        ],
        "Alta", "Dashboard"
    )

    add_story(doc,
        "HU-12", "Ver métricas de su área",
        "Directivo o Técnico",
        "ver métricas y KPIs filtrados por mi vicepresidencia o dependencia para hacer seguimiento a los hallazgos de mi área",
        [
            "Se muestran los mismos indicadores del dashboard pero filtrados por el área del usuario.",
            "Las gráficas reflejan únicamente los datos del área del usuario.",
            "El usuario no puede ver datos de otras vicepresidencias.",
        ],
        "Alta", "Dashboard"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 5: GESTIÓN DE USUARIOS
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 5: Gestión de Usuarios", 1)

    add_story(doc,
        "HU-13", "Crear nuevo usuario",
        "Vicepresidente",
        "crear nuevas cuentas de usuario con el rol y área correspondiente para dar acceso al sistema a nuevos colaboradores",
        [
            "Solo el Vicepresidente puede crear usuarios.",
            "Se requieren: nombre, correo electrónico, contraseña y rol.",
            "El correo electrónico debe ser único en el sistema.",
            "Opcionalmente se puede asignar vicepresidencia y dependencia.",
            "El usuario queda activo por defecto.",
            "Se muestra confirmación al crear exitosamente.",
        ],
        "Alta", "Usuarios"
    )

    add_story(doc,
        "HU-14", "Editar información de usuario",
        "Vicepresidente",
        "editar la información de cualquier usuario para mantener actualizada la información del sistema",
        [
            "Solo el Vicepresidente puede editar usuarios.",
            "Se pueden modificar: nombre, correo, contraseña, rol, vicepresidencia, dependencia.",
            "Al cambiar el rol, los permisos se aplican en el siguiente inicio de sesión.",
            "Se muestra confirmación al guardar exitosamente.",
        ],
        "Alta", "Usuarios"
    )

    add_story(doc,
        "HU-15", "Desactivar usuario",
        "Vicepresidente",
        "desactivar una cuenta de usuario para revocar su acceso al sistema sin eliminar su historial",
        [
            "Solo el Vicepresidente puede desactivar usuarios.",
            "Un usuario desactivado no puede iniciar sesión.",
            "El historial y registros del usuario desactivado se conservan.",
            "No se puede desactivar la propia cuenta.",
            "Se solicita confirmación antes de desactivar.",
        ],
        "Alta", "Usuarios"
    )

    add_story(doc,
        "HU-16", "Ver listado de usuarios",
        "Vicepresidente o Directivo",
        "ver el listado de usuarios del sistema para conocer quiénes tienen acceso y sus roles",
        [
            "El Vicepresidente ve todos los usuarios del sistema.",
            "El Directivo ve solo los usuarios de su área.",
            "Se muestra: nombre, correo, rol, área, estado (activo/inactivo).",
            "Se puede filtrar por rol o estado.",
        ],
        "Media", "Usuarios"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # ÉPICA 6: ACTIVIDADES
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "Épica 6: Actividades del Plan de Acción", 1)

    add_story(doc,
        "HU-17", "Ver actividades de un hallazgo",
        "usuario autenticado",
        "ver todas las actividades del plan de acción asociadas a un hallazgo para conocer el detalle de las acciones comprometidas",
        [
            "Las actividades se muestran en el detalle del hallazgo.",
            "Se muestra: ID del plan, nombre, descripción, estado, responsable, fecha de compromiso, prórroga.",
            "Múltiples actividades pueden estar vinculadas a un mismo hallazgo.",
            "Las actividades son de solo lectura (se cargan desde Excel).",
        ],
        "Media", "Actividades"
    )

    # ─────────────────────────────────────────────────────────────────────────
    # Resumen de historias
    # ─────────────────────────────────────────────────────────────────────────
    page_break(doc)
    add_heading(doc, "Resumen de Historias de Usuario", 1)
    summary = [
        ("HU-01", "Inicio de sesión con credenciales",        "Autenticación",  "Alta",  "Todos"),
        ("HU-02", "Cierre de sesión",                          "Autenticación",  "Alta",  "Todos"),
        ("HU-03", "Visualizar listado de hallazgos",           "Hallazgos",      "Alta",  "Todos"),
        ("HU-04", "Filtrar hallazgos por criterios múltiples", "Hallazgos",      "Alta",  "Todos"),
        ("HU-05", "Ver detalle de un hallazgo",                "Hallazgos",      "Alta",  "Todos"),
        ("HU-06", "Editar campos de un hallazgo",              "Hallazgos",      "Alta",  "Directivo / VP"),
        ("HU-07", "Importar hallazgos desde Excel",            "Carga Excel",    "Alta",  "Directivo / VP"),
        ("HU-08", "Analizar Excel antes de cargar",            "Carga Excel",    "Media", "Todos"),
        ("HU-09", "Ver historial de cargas",                   "Carga Excel",    "Media", "Directivo / VP"),
        ("HU-10", "Eliminar una carga y sus hallazgos",        "Carga Excel",    "Media", "Directivo / VP"),
        ("HU-11", "Ver dashboard ejecutivo global",            "Dashboard",      "Alta",  "Vicepresidente"),
        ("HU-12", "Ver métricas de su área",                   "Dashboard",      "Alta",  "Directivo / Técnico"),
        ("HU-13", "Crear nuevo usuario",                       "Usuarios",       "Alta",  "Vicepresidente"),
        ("HU-14", "Editar información de usuario",             "Usuarios",       "Alta",  "Vicepresidente"),
        ("HU-15", "Desactivar usuario",                        "Usuarios",       "Alta",  "Vicepresidente"),
        ("HU-16", "Ver listado de usuarios",                   "Usuarios",       "Media", "Directivo / VP"),
        ("HU-17", "Ver actividades de un hallazgo",            "Actividades",    "Media", "Todos"),
    ]
    cols = ["ID", "Historia", "Épica", "Prioridad", "Roles"]
    add_permission_table(doc, cols, summary)

    path = "c:/Users/Usuario/hallazgos/Historias_de_Usuario.docx"
    doc.save(path)
    print(f"[OK] Guardado: {path}")
    return path


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Generando documentación del Sistema Hallazgos ERO...")
    p1 = crear_manual_usuario()
    p2 = crear_diagrama_relaciones()
    p3 = crear_historias_usuario()
    print("\n✓ Documentación generada exitosamente:")
    print(f"  1. {p1}")
    print(f"  2. {p2}")
    print(f"  3. {p3}")
